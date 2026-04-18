import logging
import os

from docx import Document
from dotenv import load_dotenv
from sarvamai import SarvamAI

logger = logging.getLogger(__name__)

load_dotenv()

MAX_CHARS = 900


def _get_client():
    """Create and return a SarvamAI client using the environment API key."""
    api_key = os.getenv("SARVAM_API_KEY")
    if not api_key:
        raise Exception("SARVAM_API_KEY not set")
    return SarvamAI(api_subscription_key=api_key)


def _chunk_text(text: str, max_chars: int) -> list[str]:
    """
    Split *text* into a list of strings each no longer than *max_chars*.

    Splitting is done on whitespace boundaries where possible; if a single
    word exceeds *max_chars* it is split at the hard character limit.
    """
    if len(text) <= max_chars:
        return [text]

    chunks = []
    while text:
        if len(text) <= max_chars:
            chunks.append(text)
            break
        # Try to split on the last space within the limit.
        split_at = text.rfind(" ", 0, max_chars)
        if split_at == -1:
            split_at = max_chars  # No space found; hard split.
        chunks.append(text[:split_at].rstrip())
        text = text[split_at:].lstrip()
    return chunks


def translate_doc(input_path, output_path, source_lang, target_lang, mode, model="mayura:v1", speaker_gender=None, numerals_format="international", client=None):
    """
    Translate paragraphs, table cells, and section headers/footers in
    *input_path* (.docx) and write the result to *output_path* (.docx).

    Parameters
    ----------
    input_path     : str           Path to the source Word document.
    output_path    : str           Path where the translated document will be saved.
    source_lang    : str           BCP-47 language code of the source text (e.g. "en-IN").
    target_lang    : str           BCP-47 language code for the output (e.g. "hi-IN").
    mode           : str           Sarvam translation mode (e.g. "formal", "colloquial").
    model          : str           Sarvam model identifier (e.g. "mayura:v1", "sarvam-translate:v1").
    speaker_gender : str | None    Optional speaker gender ("Male" / "Female").
    numerals_format: str           Numeral style ("international" / "native").
    client         : SarvamAI | None
                                   Optional pre-built API client; one is created from the
                                   environment when None.  Pass a mock here in tests.

    Text boxes and shapes are not translated.
    """
    if client is None:
        client = _get_client()

    doc = Document(input_path)
    new_doc = Document()

    def translate_batch(texts: list[str]) -> list[str]:
        """Translate a list of non-empty strings, batching to respect MAX_CHARS.
        Returns translated strings in the same order as input."""
        if not texts:
            return []

        results = []
        buf: list[str] = []
        buf_len: int = 0

        def flush():
            if not buf:
                return
            text_block = "\n".join(buf)
            kwargs = dict(
                input=text_block,
                source_language_code=source_lang,
                target_language_code=target_lang,
                mode=mode,
                model=model,
                numerals_format=numerals_format,
            )
            if speaker_gender is not None:
                kwargs["speaker_gender"] = speaker_gender
            response = client.text.translate(**kwargs)
            result_lines = response.translated_text.split("\n")
            if len(result_lines) != len(buf):
                logger.warning(
                    "Paragraph count mismatch after translation: submitted %d, received %d.",
                    len(buf), len(result_lines),
                )
            results.extend(result_lines)
            buf.clear()

        for text in texts:
            sub_chunks = _chunk_text(text, MAX_CHARS)
            for chunk in sub_chunks:
                if buf_len + len(chunk) > MAX_CHARS:
                    flush()
                    buf_len = 0
                buf.append(chunk)
                buf_len += len(chunk)

        flush()
        return results

    # Translate paragraphs
    para_texts = []
    para_empty_flags = []
    for para in doc.paragraphs:
        text = para.text.strip()
        para_empty_flags.append(not text)
        if text:
            para_texts.append(text)

    translated_paras = translate_batch(para_texts)
    translated_iter = iter(translated_paras)
    for is_empty in para_empty_flags:
        if is_empty:
            new_doc.add_paragraph("")
        else:
            new_doc.add_paragraph(next(translated_iter, ""))

    # Translate table cells
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_texts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                if not cell_texts:
                    continue
                translated_cell_texts = translate_batch(cell_texts)
                new_cell = new_table.cell(row_idx, col_idx)
                for para in new_cell.paragraphs:
                    para.clear()
                for i, translated_text in enumerate(translated_cell_texts):
                    if i == 0:
                        new_cell.paragraphs[0].add_run(translated_text)
                    else:
                        new_cell.add_paragraph(translated_text)

    # Translate headers and footers (default section only; warn if multiple sections)
    if len(doc.sections) > 1:
        logger.warning(
            "Source document has %d sections; header/footer translation is applied to the default section only.",
            len(doc.sections),
        )

    first_section = doc.sections[0]
    new_section = new_doc.sections[0]

    header_texts = [p.text.strip() for p in first_section.header.paragraphs if p.text.strip()]
    if header_texts:
        translated_headers = translate_batch(header_texts)
        new_header_paras = new_section.header.paragraphs
        for i, translated in enumerate(translated_headers):
            if i < len(new_header_paras):
                new_header_paras[i].clear()
                new_header_paras[i].add_run(translated)

    footer_texts = [p.text.strip() for p in first_section.footer.paragraphs if p.text.strip()]
    if footer_texts:
        translated_footers = translate_batch(footer_texts)
        new_footer_paras = new_section.footer.paragraphs
        for i, translated in enumerate(translated_footers):
            if i < len(new_footer_paras):
                new_footer_paras[i].clear()
                new_footer_paras[i].add_run(translated)

    new_doc.save(output_path)
