"""
Unit tests for translate_doc.py — batching logic, chunking, and document processing.

The SarvamAI client is supplied via the `client` parameter (dependency injection),
so these tests make no network calls and require no real SARVAM_API_KEY.

Run with:
    pytest tests/test_translate_doc.py -v
"""

import os
import tempfile
from unittest.mock import MagicMock

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from translate_doc import translate_doc, _chunk_text

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx(paragraphs: list[str]) -> str:
    """Write a temporary .docx file with the given paragraphs; return its path."""
    tmpdir = tempfile.mkdtemp()
    path = os.path.join(tmpdir, "input.docx")
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(path)
    return path


def _mock_client(translated_text: str) -> MagicMock:
    """Return a mock SarvamAI client whose translate() returns *translated_text*."""
    response = MagicMock()
    response.translated_text = translated_text
    client = MagicMock()
    client.text.translate.return_value = response
    return client


# ---------------------------------------------------------------------------
# _chunk_text unit tests
# ---------------------------------------------------------------------------

class TestChunkText:

    def test_short_text_returned_as_single_chunk(self):
        assert _chunk_text("hello", 900) == ["hello"]

    def test_exact_limit_returned_as_single_chunk(self):
        text = "A" * 900
        assert _chunk_text(text, 900) == [text]

    def test_text_over_limit_is_split(self):
        text = "A" * 1800
        chunks = _chunk_text(text, 900)
        assert len(chunks) == 2
        assert all(len(c) <= 900 for c in chunks)

    def test_split_prefers_whitespace_boundary(self):
        # "hello world" is 11 chars; limit of 7 should split at the space.
        chunks = _chunk_text("hello world", 7)
        assert chunks[0] == "hello"
        assert chunks[1] == "world"

    def test_hard_split_when_no_space(self):
        text = "A" * 20
        chunks = _chunk_text(text, 10)
        assert all(len(c) <= 10 for c in chunks)
        assert "".join(chunks) == text

    def test_reconstructed_text_matches_original_whitespace_split(self):
        words = ["word"] * 300  # 1500 chars with spaces
        text = " ".join(words)
        chunks = _chunk_text(text, 900)
        assert all(len(c) <= 900 for c in chunks)
        # Joining chunks with a space should reconstruct the original.
        assert " ".join(chunks) == text


# ---------------------------------------------------------------------------
# translate_doc unit tests
# ---------------------------------------------------------------------------

class TestTranslateDocBatching:

    def test_single_short_paragraph(self, tmp_path):
        """A single short paragraph is translated and written to the output."""
        input_path = _make_docx(["Hello world"])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("नमस्ते दुनिया")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        result = Document(output_path)
        texts = [p.text for p in result.paragraphs if p.text]
        assert texts == ["नमस्ते दुनिया"]
        client.text.translate.assert_called_once()

    def test_empty_paragraphs_preserved_without_api_call(self, tmp_path):
        """Empty paragraphs are written as blank lines and never sent to the API."""
        input_path = _make_docx(["Hello", "", "World"])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("नमस्ते\nदुनिया")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        assert client.text.translate.call_count == 1
        result = Document(output_path)
        assert "" in [p.text for p in result.paragraphs]

    def test_buffer_flushes_when_limit_exceeded(self, tmp_path):
        """Two paragraphs that together exceed 900 chars cause two separate API calls."""
        input_path = _make_docx(["A" * 500, "B" * 500])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("translated")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        assert client.text.translate.call_count == 2

    def test_multiple_short_paragraphs_batched_in_one_call(self, tmp_path):
        """Several short paragraphs that fit within 900 chars go in a single API call."""
        paragraphs = [f"Line {i}" for i in range(10)]
        input_path = _make_docx(paragraphs)
        output_path = str(tmp_path / "out.docx")
        translated = "\n".join([f"पंक्ति {i}" for i in range(10)])
        client = _mock_client(translated)

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        client.text.translate.assert_called_once()

    def test_output_file_is_created(self, tmp_path):
        """The output .docx file is written to the specified path."""
        input_path = _make_docx(["Test"])
        output_path = str(tmp_path / "translated.docx")
        client = _mock_client("परीक्षण")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        assert os.path.exists(output_path)

    def test_translation_mode_forwarded(self, tmp_path):
        """The `mode` parameter is passed through to the Sarvam API call."""
        input_path = _make_docx(["Hello"])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("नमस्ते")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "colloquial", client=client)

        call_kwargs = client.text.translate.call_args.kwargs
        assert call_kwargs.get("mode") == "colloquial"

    def test_source_and_target_language_forwarded(self, tmp_path):
        """source_lang and target_lang are forwarded to the Sarvam API."""
        input_path = _make_docx(["வணக்கம்"])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("Hello")

        translate_doc(input_path, output_path, "ta-IN", "en-IN", "formal", client=client)

        call_kwargs = client.text.translate.call_args.kwargs
        assert call_kwargs.get("source_language_code") == "ta-IN"
        assert call_kwargs.get("target_language_code") == "en-IN"

    def test_all_empty_paragraphs_make_no_api_call(self, tmp_path):
        """A document with only blank paragraphs triggers no API calls."""
        input_path = _make_docx(["", "", ""])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        client.text.translate.assert_not_called()

    def test_exact_900_char_paragraph_sent_in_one_call(self, tmp_path):
        """A paragraph of exactly 900 characters is sent in a single API call."""
        input_path = _make_docx(["C" * 900])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("translated")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        client.text.translate.assert_called_once()

    def test_oversized_single_paragraph_is_chunked(self, tmp_path):
        """BUG-2 fix: a paragraph > 900 chars is split and each chunk sent separately."""
        long_para = "word " * 250  # ~1250 chars; must be chunked
        input_path = _make_docx([long_para.strip()])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("translated")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        # Must have made more than one API call because the para exceeds 900 chars.
        assert client.text.translate.call_count > 1
        # Every submitted chunk must be <= 900 chars.
        for call in client.text.translate.call_args_list:
            submitted = call.kwargs.get("input", "")
            assert len(submitted) <= 900

    def test_paragraph_count_mismatch_logs_warning(self, tmp_path, caplog):
        """BUG-3 fix: a mismatch between submitted and returned paragraph counts is logged."""
        import logging
        input_path = _make_docx(["Line one", "Line two"])
        output_path = str(tmp_path / "out.docx")
        # API returns only one line instead of two → mismatch.
        client = _mock_client("Only one line returned")

        with caplog.at_level(logging.WARNING, logger="translate_doc"):
            translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        assert any("mismatch" in record.message.lower() for record in caplog.records)

    def test_speaker_gender_forwarded_when_provided(self, tmp_path):
        """speaker_gender is included in the API call when provided."""
        input_path = _make_docx(["Hello"])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("नमस्ते")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", speaker_gender="Male", client=client)

        call_kwargs = client.text.translate.call_args.kwargs
        assert call_kwargs.get("speaker_gender") == "Male"

    def test_speaker_gender_omitted_when_none(self, tmp_path):
        """speaker_gender is not included in the API call when None."""
        input_path = _make_docx(["Hello"])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("नमस्ते")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", speaker_gender=None, client=client)

        call_kwargs = client.text.translate.call_args.kwargs
        assert "speaker_gender" not in call_kwargs

    def test_numerals_format_forwarded(self, tmp_path):
        """numerals_format is always included in the API call."""
        input_path = _make_docx(["Hello"])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("नमस्ते")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", numerals_format="native", client=client)

        call_kwargs = client.text.translate.call_args.kwargs
        assert call_kwargs.get("numerals_format") == "native"

    def test_numerals_format_default_is_international(self, tmp_path):
        """numerals_format defaults to 'international' when not specified."""
        input_path = _make_docx(["Hello"])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("नमस्ते")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        call_kwargs = client.text.translate.call_args.kwargs
        assert call_kwargs.get("numerals_format") == "international"

    def test_model_forwarded_to_api_call(self, tmp_path):
        """model is included in every Sarvam API call."""
        input_path = _make_docx(["Hello"])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("नमस्ते")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", model="sarvam-translate:v1", client=client)

        call_kwargs = client.text.translate.call_args.kwargs
        assert call_kwargs.get("model") == "sarvam-translate:v1"

    def test_default_model_is_mayura(self, tmp_path):
        """model defaults to 'mayura:v1' when not specified."""
        input_path = _make_docx(["Hello"])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("नमस्ते")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        call_kwargs = client.text.translate.call_args.kwargs
        assert call_kwargs.get("model") == "mayura:v1"

    def test_no_api_key_raises_without_client(self, tmp_path, monkeypatch):
        """When no client is provided and SARVAM_API_KEY is absent, an exception is raised."""
        monkeypatch.delenv("SARVAM_API_KEY", raising=False)
        input_path = _make_docx(["Hello"])
        output_path = str(tmp_path / "out.docx")

        with pytest.raises(Exception, match="SARVAM_API_KEY not set"):
            translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal")

    def test_retry_succeeds_after_two_failures(self, tmp_path):
        """Transient failures on the first two calls are retried; third attempt succeeds."""
        response = MagicMock()
        response.translated_text = "नमस्ते"
        client = MagicMock()
        client.text.translate.side_effect = [
            Exception("transient error"),
            Exception("transient error"),
            response,
        ]

        input_path = _make_docx(["Hello"])
        output_path = str(tmp_path / "out.docx")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        assert client.text.translate.call_count == 3
        result = Document(output_path)
        assert any(p.text == "नमस्ते" for p in result.paragraphs)

    def test_retry_reraises_after_three_failures(self, tmp_path):
        """After three consecutive failures the original exception propagates."""
        client = MagicMock()
        client.text.translate.side_effect = Exception("permanent error")

        input_path = _make_docx(["Hello"])
        output_path = str(tmp_path / "out.docx")

        with pytest.raises(Exception, match="permanent error"):
            translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)


# ---------------------------------------------------------------------------
# Document coverage tests — tables, headers, footers
# ---------------------------------------------------------------------------

def _make_docx_with_table(cell_texts: list[list[str]]) -> str:
    """Create a temporary .docx with a single table; return its path."""
    tmpdir = tempfile.mkdtemp()
    path = os.path.join(tmpdir, "input.docx")
    doc = Document()
    rows = len(cell_texts)
    cols = max(len(row) for row in cell_texts) if cell_texts else 1
    table = doc.add_table(rows=rows, cols=cols)
    for r, row in enumerate(cell_texts):
        for c, text in enumerate(row):
            if text:
                table.cell(r, c).paragraphs[0].add_run(text)
    doc.save(path)
    return path


def _make_docx_with_header(header_text: str) -> str:
    """Create a temporary .docx with a header paragraph; return its path."""
    tmpdir = tempfile.mkdtemp()
    path = os.path.join(tmpdir, "input.docx")
    doc = Document()
    doc.sections[0].header.paragraphs[0].add_run(header_text)
    doc.save(path)
    return path


def _make_docx_with_footer(footer_text: str) -> str:
    """Create a temporary .docx with a footer paragraph; return its path."""
    tmpdir = tempfile.mkdtemp()
    path = os.path.join(tmpdir, "input.docx")
    doc = Document()
    doc.sections[0].footer.paragraphs[0].add_run(footer_text)
    doc.save(path)
    return path


class TestDocumentCoverage:

    def test_table_cell_text_is_translated(self, tmp_path):
        """Non-empty table cell text is passed to the Sarvam API."""
        input_path = _make_docx_with_table([["Hello", "World"]])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("Translated")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        all_inputs = [call.kwargs.get("input", "") for call in client.text.translate.call_args_list]
        assert any("Hello" in inp for inp in all_inputs)

    def test_empty_table_cells_not_translated(self, tmp_path):
        """Empty table cells do not trigger an API call."""
        input_path = _make_docx_with_table([["", ""]])
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        client.text.translate.assert_not_called()

    def test_header_text_is_translated(self, tmp_path):
        """Text in the document header is passed to the Sarvam API."""
        input_path = _make_docx_with_header("Page Header")
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("Translated Header")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        all_inputs = [call.kwargs.get("input", "") for call in client.text.translate.call_args_list]
        assert any("Page Header" in inp for inp in all_inputs)

    def test_footer_text_is_translated(self, tmp_path):
        """Text in the document footer is passed to the Sarvam API."""
        input_path = _make_docx_with_footer("Page Footer")
        output_path = str(tmp_path / "out.docx")
        client = _mock_client("Translated Footer")

        translate_doc(input_path, output_path, "en-IN", "hi-IN", "formal", client=client)

        all_inputs = [call.kwargs.get("input", "") for call in client.text.translate.call_args_list]
        assert any("Page Footer" in inp for inp in all_inputs)


# ---------------------------------------------------------------------------
# Paragraph formatting preservation tests
# ---------------------------------------------------------------------------

class TestParagraphFormatting:

    def test_heading1_style_preserved(self, tmp_path):
        """A paragraph styled as 'Heading 1' retains that style in the output."""
        tmpdir = tempfile.mkdtemp()
        path = os.path.join(tmpdir, "input.docx")
        doc = Document()
        doc.add_paragraph("Hello", style="Heading 1")
        doc.save(path)

        output_path = str(tmp_path / "out.docx")
        client = _mock_client("नमस्ते")

        translate_doc(path, output_path, "en-IN", "hi-IN", "formal", client=client)

        result = Document(output_path)
        non_empty = [p for p in result.paragraphs if p.text.strip()]
        assert non_empty[0].style.name == "Heading 1"

    def test_center_alignment_preserved(self, tmp_path):
        """A centered paragraph retains WD_ALIGN_PARAGRAPH.CENTER in the output."""
        tmpdir = tempfile.mkdtemp()
        path = os.path.join(tmpdir, "input.docx")
        doc = Document()
        para = doc.add_paragraph("Centered text")
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(path)

        output_path = str(tmp_path / "out.docx")
        client = _mock_client("केंद्रित पाठ")

        translate_doc(path, output_path, "en-IN", "hi-IN", "formal", client=client)

        result = Document(output_path)
        non_empty = [p for p in result.paragraphs if p.text.strip()]
        assert non_empty[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_unknown_style_does_not_raise(self, tmp_path):
        """A source paragraph with a custom style absent from a fresh Document does not raise."""
        from docx.enum.style import WD_STYLE_TYPE
        tmpdir = tempfile.mkdtemp()
        path = os.path.join(tmpdir, "input.docx")
        doc = Document()
        doc.styles.add_style("MyCustomStyle", WD_STYLE_TYPE.PARAGRAPH)
        para = doc.add_paragraph("Hello")
        para.style = doc.styles["MyCustomStyle"]
        doc.save(path)

        output_path = str(tmp_path / "out.docx")
        client = _mock_client("नमस्ते")

        translate_doc(path, output_path, "en-IN", "hi-IN", "formal", client=client)

        assert os.path.exists(output_path)
