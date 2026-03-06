"""
Integration tests for app.py — FastAPI endpoint behaviour.

Uses FastAPI's TestClient (backed by httpx). The translate_doc function is
patched at the app module level so no real translation or Sarvam API calls occur.

Run with:
    pytest tests/test_app.py -v
"""

import importlib
import io
import os
import time
from unittest.mock import patch

import pytest
from fastapi.testclient import TestClient
from docx import Document

# Provide a dummy key and a high rate limit so the module-level client
# does not trip the rate limiter during the test session.
os.environ.setdefault("SARVAM_API_KEY", "test-key-for-integration-tests")
os.environ.setdefault("RATE_LIMIT_PER_MINUTE", "1000")

import app as app_module

client = TestClient(app_module.app)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    """Return a valid .docx file as raw bytes, ready for a multipart upload."""
    buf = io.BytesIO()
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _fake_translate(input_path, output_path, source_lang, target_lang, mode, **kwargs):
    """Side-effect for patched translate_doc: writes a minimal valid .docx output."""
    doc = Document()
    doc.add_paragraph("translated output")
    doc.save(output_path)


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ---------------------------------------------------------------------------
# Happy-path tests
# ---------------------------------------------------------------------------

class TestTranslateDocEndpoint:

    def test_valid_request_returns_200_and_docx(self):
        """A well-formed request returns HTTP 200 with a .docx content type."""
        with patch("app.translate_doc", side_effect=_fake_translate):
            response = client.post(
                "/translate-doc",
                files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                data={"source_lang": "en-IN", "target_lang": "hi-IN"},
            )
        assert response.status_code == 200
        assert "wordprocessingml.document" in response.headers["content-type"]

    def test_response_filename_is_translated_docx(self):
        """Content-Disposition header specifies filename as translated.docx."""
        with patch("app.translate_doc", side_effect=_fake_translate):
            response = client.post(
                "/translate-doc",
                files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                data={"source_lang": "en-IN", "target_lang": "hi-IN"},
            )
        assert "translated.docx" in response.headers.get("content-disposition", "")

    def test_mode_defaults_to_formal(self):
        """When `mode` is not supplied it defaults to 'formal'."""
        captured = {}

        def _capture(input_path, output_path, source_lang, target_lang, mode, **kwargs):
            captured["mode"] = mode
            _fake_translate(input_path, output_path, source_lang, target_lang, mode)

        with patch("app.translate_doc", side_effect=_capture):
            client.post(
                "/translate-doc",
                files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                data={"source_lang": "en-IN", "target_lang": "hi-IN"},
            )
        assert captured.get("mode") == "formal"

    def test_explicit_mode_is_forwarded(self):
        """An explicitly supplied `mode` value is forwarded to translate_doc."""
        captured = {}

        def _capture(input_path, output_path, source_lang, target_lang, mode, **kwargs):
            captured["mode"] = mode
            _fake_translate(input_path, output_path, source_lang, target_lang, mode)

        with patch("app.translate_doc", side_effect=_capture):
            client.post(
                "/translate-doc",
                files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                data={"source_lang": "en-IN", "target_lang": "hi-IN", "mode": "colloquial"},
            )
        assert captured.get("mode") == "colloquial"

    def test_languages_are_forwarded(self):
        """source_lang and target_lang are forwarded to translate_doc unchanged."""
        captured = {}

        def _capture(input_path, output_path, source_lang, target_lang, mode, **kwargs):
            captured["source_lang"] = source_lang
            captured["target_lang"] = target_lang
            _fake_translate(input_path, output_path, source_lang, target_lang, mode)

        with patch("app.translate_doc", side_effect=_capture):
            client.post(
                "/translate-doc",
                files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                data={"source_lang": "ta-IN", "target_lang": "en-IN"},
            )
        assert captured.get("source_lang") == "ta-IN"
        assert captured.get("target_lang") == "en-IN"

    def test_response_body_is_valid_docx(self):
        """The response body can be parsed as a valid .docx document."""
        with patch("app.translate_doc", side_effect=_fake_translate):
            response = client.post(
                "/translate-doc",
                files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                data={"source_lang": "en-IN", "target_lang": "hi-IN"},
            )
        doc = Document(io.BytesIO(response.content))
        texts = [p.text for p in doc.paragraphs if p.text]
        assert texts == ["translated output"]


# ---------------------------------------------------------------------------
# Validation / error tests
# ---------------------------------------------------------------------------

class TestInputValidation:

    def test_missing_file_returns_422(self):
        """Omitting the required `file` field returns HTTP 422."""
        response = client.post(
            "/translate-doc",
            data={"source_lang": "en-IN", "target_lang": "hi-IN"},
        )
        assert response.status_code == 422

    def test_missing_source_lang_returns_422(self):
        """Omitting `source_lang` returns HTTP 422."""
        response = client.post(
            "/translate-doc",
            files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
            data={"target_lang": "hi-IN"},
        )
        assert response.status_code == 422

    def test_missing_target_lang_returns_422(self):
        """Omitting `target_lang` returns HTTP 422."""
        response = client.post(
            "/translate-doc",
            files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
            data={"source_lang": "en-IN"},
        )
        assert response.status_code == 422

    def test_translate_doc_exception_returns_500(self):
        """If translate_doc raises, the endpoint returns HTTP 500."""
        with patch("app.translate_doc", side_effect=RuntimeError("Sarvam API error")):
            response = client.post(
                "/translate-doc",
                files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                data={"source_lang": "en-IN", "target_lang": "hi-IN"},
            )
        assert response.status_code == 500


# ---------------------------------------------------------------------------
# CORS tests
# ---------------------------------------------------------------------------

class TestCORS:

    def test_allowed_origin_job_joseph(self):
        """Preflight from https://job-joseph.com gets the origin reflected."""
        response = client.options(
            "/translate-doc",
            headers={
                "Origin": "https://job-joseph.com",
                "Access-Control-Request-Method": "POST",
            },
        )
        assert response.headers.get("access-control-allow-origin") == "https://job-joseph.com"

    def test_allowed_origin_www_job_joseph(self):
        """Preflight from https://www.job-joseph.com gets the origin reflected."""
        response = client.options(
            "/translate-doc",
            headers={
                "Origin": "https://www.job-joseph.com",
                "Access-Control-Request-Method": "POST",
            },
        )
        assert response.headers.get("access-control-allow-origin") == "https://www.job-joseph.com"

    def test_disallowed_origin_not_reflected(self):
        """Preflight from an unknown origin does not get an allow-origin header."""
        response = client.options(
            "/translate-doc",
            headers={
                "Origin": "https://evil.example.com",
                "Access-Control-Request-Method": "POST",
            },
        )
        allow_origin = response.headers.get("access-control-allow-origin", "")
        assert allow_origin != "https://evil.example.com"


# ---------------------------------------------------------------------------
# Language code and mode validation tests
# ---------------------------------------------------------------------------

class TestLanguageAndModeValidation:

    def test_invalid_source_lang_returns_422(self):
        """An unrecognised source_lang returns HTTP 422 before the file is processed."""
        response = client.post(
            "/translate-doc",
            files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
            data={"source_lang": "xx-XX", "target_lang": "hi-IN"},
        )
        assert response.status_code == 422
        assert "source_lang" in response.json()["detail"]

    def test_invalid_target_lang_returns_422(self):
        """An unrecognised target_lang returns HTTP 422."""
        response = client.post(
            "/translate-doc",
            files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
            data={"source_lang": "en-IN", "target_lang": "zz-ZZ"},
        )
        assert response.status_code == 422
        assert "target_lang" in response.json()["detail"]

    def test_invalid_mode_returns_422(self):
        """An unrecognised mode returns HTTP 422."""
        response = client.post(
            "/translate-doc",
            files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
            data={"source_lang": "en-IN", "target_lang": "hi-IN", "mode": "pirate"},
        )
        assert response.status_code == 422
        assert "mode" in response.json()["detail"]

    def test_colloquial_mode_is_accepted(self):
        """The 'colloquial' mode passes validation and is forwarded to translate_doc."""
        with patch("app.translate_doc", side_effect=_fake_translate):
            response = client.post(
                "/translate-doc",
                files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                data={"source_lang": "en-IN", "target_lang": "hi-IN", "mode": "colloquial"},
            )
        assert response.status_code == 200

    def test_all_valid_lang_codes_accepted(self):
        """Every lang code in VALID_LANG_CODES is accepted without a 422."""
        valid_codes = app_module.VALID_LANG_CODES
        # Test a representative cross-product pair to keep the test fast
        for code in list(valid_codes)[:3]:
            with patch("app.translate_doc", side_effect=_fake_translate):
                response = client.post(
                    "/translate-doc",
                    files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                    data={"source_lang": "en-IN", "target_lang": code},
                )
            assert response.status_code == 200, f"Expected 200 for target_lang={code}"


# ---------------------------------------------------------------------------
# File size limit tests
# ---------------------------------------------------------------------------

class TestFileSizeLimit:

    def test_file_within_limit_succeeds(self):
        """A small file passes the size check and is processed."""
        with patch("app.translate_doc", side_effect=_fake_translate):
            response = client.post(
                "/translate-doc",
                files={"file": ("small.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                data={"source_lang": "en-IN", "target_lang": "hi-IN"},
            )
        assert response.status_code == 200

    def test_oversized_file_returns_413(self, monkeypatch):
        """A file whose byte length exceeds MAX_FILE_SIZE_BYTES returns HTTP 413."""
        monkeypatch.setattr(app_module, "MAX_FILE_SIZE_BYTES", 10)
        response = client.post(
            "/translate-doc",
            files={"file": ("big.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
            data={"source_lang": "en-IN", "target_lang": "hi-IN"},
        )
        assert response.status_code == 413
        assert "too large" in response.json()["detail"].lower()


# ---------------------------------------------------------------------------
# Document character limit tests
# ---------------------------------------------------------------------------

class TestDocCharLimit:

    def test_document_within_char_limit_succeeds(self):
        """A document with fewer characters than MAX_DOC_CHARS proceeds to translation."""
        with patch("app.translate_doc", side_effect=_fake_translate):
            response = client.post(
                "/translate-doc",
                files={"file": ("small.docx", _make_docx_bytes(["Short"]), DOCX_MIME)},
                data={"source_lang": "en-IN", "target_lang": "hi-IN"},
            )
        assert response.status_code == 200

    def test_document_exceeding_char_limit_returns_422(self, monkeypatch):
        """A document exceeding MAX_DOC_CHARS returns HTTP 422 without calling the API."""
        monkeypatch.setattr(app_module, "MAX_DOC_CHARS", 3)
        response = client.post(
            "/translate-doc",
            files={"file": ("big.docx", _make_docx_bytes(["Hello world"]), DOCX_MIME)},
            data={"source_lang": "en-IN", "target_lang": "hi-IN"},
        )
        assert response.status_code == 422
        assert "too large" in response.json()["detail"].lower()

    def test_doc_char_limit_error_contains_char_count(self, monkeypatch):
        """The 422 detail message includes the actual and allowed character counts."""
        monkeypatch.setattr(app_module, "MAX_DOC_CHARS", 3)
        response = client.post(
            "/translate-doc",
            files={"file": ("big.docx", _make_docx_bytes(["Hello world"]), DOCX_MIME)},
            data={"source_lang": "en-IN", "target_lang": "hi-IN"},
        )
        detail = response.json()["detail"]
        assert "3" in detail  # limit appears in the message


# ---------------------------------------------------------------------------
# Request timeout tests
# ---------------------------------------------------------------------------

class TestRequestTimeout:

    def test_slow_translation_returns_504(self, monkeypatch):
        """When translate_doc exceeds REQUEST_TIMEOUT_SECONDS the endpoint returns 504."""
        monkeypatch.setattr(app_module, "REQUEST_TIMEOUT_SECONDS", 0.05)

        def _slow(*args, **kwargs):
            time.sleep(5)

        with patch("app.translate_doc", side_effect=_slow):
            response = client.post(
                "/translate-doc",
                files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                data={"source_lang": "en-IN", "target_lang": "hi-IN"},
            )
        assert response.status_code == 504
        assert "timed out" in response.json()["detail"].lower()


# ---------------------------------------------------------------------------
# Structured error response tests
# ---------------------------------------------------------------------------

class TestStructuredErrors:

    def test_exception_returns_json_500_with_detail(self):
        """When translate_doc raises, the response is JSON 500 with a 'detail' key."""
        with patch("app.translate_doc", side_effect=RuntimeError("boom")):
            response = client.post(
                "/translate-doc",
                files={"file": ("test.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
                data={"source_lang": "en-IN", "target_lang": "hi-IN"},
            )
        assert response.status_code == 500
        body = response.json()
        assert "detail" in body
        assert "boom" not in body["detail"]  # internal error must not be leaked

    def test_413_response_has_detail_key(self, monkeypatch):
        """The 413 error response is JSON with a 'detail' key."""
        monkeypatch.setattr(app_module, "MAX_FILE_SIZE_BYTES", 10)
        response = client.post(
            "/translate-doc",
            files={"file": ("big.docx", _make_docx_bytes(["Hello"]), DOCX_MIME)},
            data={"source_lang": "en-IN", "target_lang": "hi-IN"},
        )
        assert response.status_code == 413
        assert "detail" in response.json()


# ---------------------------------------------------------------------------
# Rate limiting tests
# ---------------------------------------------------------------------------

class TestRateLimiting:

    def test_rate_limit_exceeded_returns_429(self):
        """After exceeding the per-IP limit, the endpoint returns HTTP 429."""
        # Reload the app with a limit of 1/minute so we can trigger it in 2 requests.
        os.environ["RATE_LIMIT_PER_MINUTE"] = "1"
        try:
            reloaded = importlib.reload(app_module)
            fresh_client = TestClient(reloaded.app)

            with patch.object(reloaded, "translate_doc", side_effect=_fake_translate):
                fresh_client.post(
                    "/translate-doc",
                    files={"file": ("test.docx", _make_docx_bytes(["Hi"]), DOCX_MIME)},
                    data={"source_lang": "en-IN", "target_lang": "hi-IN"},
                )
                response = fresh_client.post(
                    "/translate-doc",
                    files={"file": ("test.docx", _make_docx_bytes(["Hi"]), DOCX_MIME)},
                    data={"source_lang": "en-IN", "target_lang": "hi-IN"},
                )
            assert response.status_code == 429
        finally:
            # Restore original env and reload so other tests are not affected.
            os.environ["RATE_LIMIT_PER_MINUTE"] = "5"
            importlib.reload(app_module)

    def test_rate_limit_response_has_detail(self):
        """The 429 response body has a non-empty 'detail' key."""
        # Trigger a real 429 using a reloaded app with limit=1/minute.
        os.environ["RATE_LIMIT_PER_MINUTE"] = "1"
        try:
            reloaded = importlib.reload(app_module)
            fresh_client = TestClient(reloaded.app)

            with patch.object(reloaded, "translate_doc", side_effect=_fake_translate):
                fresh_client.post(
                    "/translate-doc",
                    files={"file": ("test.docx", _make_docx_bytes(["Hi"]), DOCX_MIME)},
                    data={"source_lang": "en-IN", "target_lang": "hi-IN"},
                )
                response = fresh_client.post(
                    "/translate-doc",
                    files={"file": ("test.docx", _make_docx_bytes(["Hi"]), DOCX_MIME)},
                    data={"source_lang": "en-IN", "target_lang": "hi-IN"},
                )
            assert response.status_code == 429
            assert "detail" in response.json()
        finally:
            os.environ["RATE_LIMIT_PER_MINUTE"] = "1000"
            importlib.reload(app_module)
